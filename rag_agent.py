import os
import logging
from typing import List
from dataclasses import dataclass
from dotenv import load_dotenv
import PyPDF2
import docx

from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain.schema import Document
from langchain_openai import ChatOpenAI

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
openai_api_key = "Your-api-key"

# Constants for text splitting
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

@dataclass
class Config:
    OPENAI_API_KEY: str
    LLM_MODEL: str = "gpt-4"
    TEMPERATURE: float = 0

@dataclass
class QueryResult:
    content: str
    confidence: float
    source: str
    metadata: dict

def load_documents(docs_directory: str) -> List[Document]:
    """Load documents from a directory, supporting .pdf and .docx formats."""
    docs = []
    for filename in os.listdir(docs_directory):
        file_path = os.path.join(docs_directory, filename)
        if filename.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = ""
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content += text
                doc_obj = Document(page_content=content, metadata={'source': filename, 'id': filename})
                docs.append(doc_obj)
        elif filename.endswith('.docx'):
            doc_file = docx.Document(file_path)
            content = ""
            for paragraph in doc_file.paragraphs:
                content += paragraph.text
            doc_obj = Document(page_content=content, metadata={'source': filename, 'id': filename})
            docs.append(doc_obj)
    return docs

class CascadeRetriever:
    """
    Implements a cascade retrieval approach:
      1. Perform sparse retrieval (BM25) on the full corpus.
      2. On the candidate documents returned, perform dense retrieval (FAISS) to rerank.
    """
    def __init__(self, documents: List[Document], openai_api_key: str):
        self.documents = documents
        self.openai_api_key = openai_api_key

    def retrieve(self, query: str, sparse_k: int = 10, dense_k: int = 5) -> List[Document]:
        # Step 1: Sparse retrieval using BM25
        sparse_retriever = BM25Retriever.from_documents(self.documents, k=sparse_k)
        candidate_docs = sparse_retriever.invoke(query)
        logger.info(f"Sparse retrieval found {len(candidate_docs)} candidate documents.")
        
        if not candidate_docs:
            return []
        # performing the embeddings on retrieved docs by the sparse
        embeddings = OpenAIEmbeddings(api_key=self.openai_api_key)
        dense_store_obj = FAISS.from_documents(candidate_docs, embeddings)
        if isinstance(dense_store_obj, tuple):
            dense_store = dense_store_obj[0]
        else:
            dense_store = dense_store_obj
        dense_retriever = dense_store.as_retriever(search_kwargs={"k": dense_k})
        final_docs = dense_retriever.get_relevant_documents(query)
        logger.info(f"Dense retrieval on candidates returned {len(final_docs)} documents.")
        return final_docs

def run_rag_agent(query: str, docs_directory: str) -> QueryResult:
    # Load all documents from the directory
    documents = load_documents(docs_directory)
    
    # Create configuration object
    config = Config(OPENAI_API_KEY=openai_api_key)
    
    # Initialize the LLM
    llm = ChatOpenAI(model=config.LLM_MODEL, temperature=config.TEMPERATURE, api_key=config.OPENAI_API_KEY)
    
    # Cascade retrieval: first sparse then dense
    cascade_retriever = CascadeRetriever(documents, config.OPENAI_API_KEY)
    final_docs = cascade_retriever.retrieve(query)
    
    if not final_docs:
        return QueryResult(
            content="No relevant information found.",
            confidence=0.0,
            source="cascade",
            metadata={"found_docs": 0}
        )
    
    # Combine retrieved documents (use only the first 500 characters of each to avoid token overload)
    combined_context = "\n".join(doc.page_content[:500] for doc in final_docs)
    
    # Create prompt for the LLM
    prompt = (
        f"Based on the following context, answer the question: {query}\n\n"
        f"Context: {combined_context}"
    )
    response = llm.invoke(prompt)
    return QueryResult(
        content=response.content,
        confidence=0.8,
        source="cascade",
        metadata={"found_docs": len(final_docs)}
    )

if __name__ == "__main__":
    docs_directory = "/Users/tharun/Desktop/Research_papers"  # Replace with your documents directory
    query = "Who is Tharun reddy Pyayala?"
    
    result = run_rag_agent(query, docs_directory)
    
    print("LLM Response:")
    print(result.content)
